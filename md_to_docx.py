import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_docx():
    doc = Document()
    
    # Title
    title = doc.add_heading('暨大选课助手 (JNU Course Grabber) 使用文档', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph('欢迎使用本选课助手！这是一款基于本地浏览器环境与 Python API 构建的高级全自动化抢课工具。它内置了智能嗅探、跨库检索、时间表防冲突预警等功能，不仅能帮助你以毫秒级速度“盲狙”抢课，还能为你提供美观、直观的图形化控制面板。')
    
    # Section 1
    doc.add_heading('核心功能亮点', level=1)
    doc.add_paragraph('1. 零配置自动登录与环境接管：软件会拉起一个专属浏览器，你只需要像平常一样登录教务系统，后台会自动嗅探所有必要的认证 Token。')
    doc.add_paragraph('2. 跨库隔离检索系统：突破教务处的“批次码隔离”限制！支持在网页端随意停留的情况下，通过高级设置中的“专业课搜索库”和“通选课搜索库”进行跨池子靶向检索。')
    doc.add_paragraph('3. 完全无感的后台静默嗅探：当你在教务系统中点击“推荐班”、“方案内”等专业课页签时，系统会自动在后台将显示的课程同步到你的待抢目标列表中。')
    doc.add_paragraph('4. 强隔离的通选课保护机制：为防止通选课列表爆炸，通选课（XGXK）不会被自动吸入。必须由你主动在“全校课程检索”里搜索并点击“加入目标”来精准入库。')
    doc.add_paragraph('5. 智能排课冲突预警：当检测到待选课程与你已选的课程在时间上发生碰撞时，目标列表中会亮起醒目的 ⚠️ 冲突 红色警报，并显示具体冲突原因。')
    doc.add_paragraph('6. 动态网格课表联动：自带实时可视化的课程表界面，选中状态一目了然。')
    doc.add_paragraph('7. 本地离线持久化：你勾选的抢课目标会被保存在本地，软件重启后可无缝恢复。')
    
    # Section 2
    doc.add_heading('安装与启动指南', level=1)
    doc.add_heading('1. 环境准备', level=2)
    doc.add_paragraph('你需要有一台装有 Python 环境的 Windows 电脑。并且需要安装少量的依赖库：\n- 打开文件夹中的 requirements.txt\n- 若未安装过所需库，请运行命令：pip install -r requirements.txt')
    doc.add_heading('2. 启动程序', level=2)
    doc.add_paragraph('非常简单，只需双击文件夹内的 "抢课系统，启动！.bat" 文件。\n- 它会自动启动 Python 后端。\n- 随后会自动弹出一个操作控制面板（网页），以及一个用于登录教务处的 Playwright 自动化浏览器。')
    
    # Section 3
    doc.add_heading('核心操作流 (How to Use)', level=1)
    
    doc.add_heading('第一步：完成教务系统接入', level=2)
    doc.add_paragraph('1. 在弹出的浏览器窗口中，输入你的账号密码登录暨大教务系统。')
    doc.add_paragraph('2. 进入到“学生选课”界面。')
    doc.add_paragraph('3. 随意点击几个界面（例如点击进入一下“专业课”轮次），控制面板右上角的终端会提示“系统接入成功，已捕获学号和 Token”。')
    
    doc.add_heading('第二步：选择你的攻击目标 (挑选课程)', level=2)
    doc.add_paragraph('你有两种方式将心仪的课程加入“待抢列表”：')
    
    doc.add_heading('方式 A：被动嗅探 (适用于专业课)', level=3)
    doc.add_paragraph('直接在教务系统中点击你的“方案内课程”或“推荐班课程”页签。控制台会自动弹出“拦截成功”的提示，同时这些课程会立刻出现在前端控制台的左侧目标列表中。')
    
    doc.add_heading('方式 B：跨库定向搜索 (适用于通选课/特定选修课)', level=3)
    doc.add_paragraph('若你想抢通识选修课，或者想打破限制搜索特定的课：\n1. 在界面的 "全校课程检索" 模块中，选择你的校区。\n2. 旁边的下拉框选择对应的底层池子：通选课搜索库 或是 专业课搜索库。\n3. 输入关键词（如：体育、音乐 等）并点击检索。\n4. 看到心仪的课后，点击右侧表格中的 "加入目标"。这门课就会立刻带着正确的属性（XGXK 或是 FANKC）加入到左侧抢课队列中。')
    
    doc.add_heading('第三步：火力全开 (开始抢课)', level=2)
    doc.add_paragraph('1. 在左侧的“目标选择”列表中，勾选那些你想抢的课程前面的复选框。（勾选状态会自动保存至本地，下次打开还在）')
    doc.add_paragraph('2. 如果课程时间有冲突，下方会显示 ⚠️ 红色警告（系统也会智能提醒）。')
    doc.add_paragraph('3. 设定合适的“请求间隔(秒)”，一般建议 0.5 ~ 1.0 秒防止被教务系统封禁。')
    doc.add_paragraph('4. 点击 “开始抢课”（或设置定时抢课）。')
    doc.add_paragraph('5. 成功后，终端会有绿色高亮提示，并且自动更新界面。')
    
    # Section 4
    doc.add_heading('注意事项与防封指南', level=1)
    doc.add_paragraph('- 不要把请求间隔设置得极低：虽然教务系统可能没有强力的频率限制，但发送太快可能会导致账号被临时屏蔽，建议使用默认的 1.0 秒，或者最低不低于 0.5 秒。')
    doc.add_paragraph('- 抢课期间，请勿在被控浏览器里随意乱点，以免打断后台的鉴权数据同步。')
    doc.add_paragraph('- 若发现软件异常或一直显示“未登录”，请刷新被控浏览器页面，或直接重启本软件。')
    
    doc.add_paragraph('祝新学期选课顺利，欧气满满！🎉')
    
    doc.save('C:/Users/a1028/Documents/antigravity/radiant-bohr/使用文档.docx')
    print("Docx created successfully.")

if __name__ == '__main__':
    create_docx()
